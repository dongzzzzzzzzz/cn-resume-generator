#!/usr/bin/env python3
"""Render a Chinese resume package from structured JSON.

Expected JSON shape is intentionally flexible. Recommended fields:

{
  "candidate": {
    "name": "候选人A",
    "target_role": "产品经理",
    "phone": "138-0000-0000",
    "email": "demo@example.com",
    "location": "上海",
    "gender": "女",
    "birth": "1998.06",
    "years_of_experience": "3年经验",
    "summary": ["3年产品经验...", "..."]
  },
  "sections": [
    {
      "title": "工作经历",
      "items": [
        {
          "heading": "某科技公司｜产品经理",
          "subheading": "增长产品组",
          "dates": "2022.07-至今",
          "bullets": ["负责...", "推动..."]
        }
      ]
    }
  ]
}
"""

from __future__ import annotations

import argparse
import json
from pathlib import Path
from typing import Any


def _text(value: Any) -> str:
    if value is None:
        return ""
    if isinstance(value, str):
        return value.strip()
    return str(value).strip()


def _as_list(value: Any) -> list[Any]:
    if value is None:
        return []
    if isinstance(value, list):
        return value
    return [value]


def _candidate(data: dict[str, Any]) -> dict[str, Any]:
    return data.get("candidate") or data.get("profile") or {}


def _derive_sections(data: dict[str, Any]) -> list[dict[str, Any]]:
    cand = _candidate(data)
    summary = cand.get("summary") or data.get("summary")
    if isinstance(data.get("sections"), list):
        sections = list(data["sections"])
        titles = {_text(section.get("title")) for section in sections if isinstance(section, dict)}
        if summary and "个人优势" not in titles:
            sections.insert(0, {"title": "个人优势", "items": _as_list(summary)})
        return sections

    mapping = [
        ("工作经历", "experience"),
        ("项目经历", "projects"),
        ("教育经历", "education"),
        ("技能证书", "skills"),
        ("奖项/作品", "awards"),
    ]
    sections: list[dict[str, Any]] = []
    if summary:
        sections.append({"title": "个人优势", "items": _as_list(summary)})
    for title, key in mapping:
        value = data.get(key)
        if value:
            sections.append({"title": title, "items": _as_list(value)})
    return sections


def _contact_parts(cand: dict[str, Any]) -> list[str]:
    return [
        _text(cand.get("phone")),
        _text(cand.get("email")),
        _text(cand.get("gender")),
        _text(cand.get("birth") or cand.get("birthday")),
        _text(cand.get("years_of_experience") or cand.get("experience_years")),
        _text(cand.get("location")),
        _text(cand.get("portfolio") or cand.get("link")),
    ]


def _item_lines(item: Any) -> tuple[str, str, str, list[str]]:
    if isinstance(item, str):
        return "", "", "", [item]
    if not isinstance(item, dict):
        return "", "", "", [_text(item)]
    heading = _text(item.get("heading") or item.get("title") or item.get("company") or item.get("school"))
    subheading = _text(item.get("subheading") or item.get("role") or item.get("major"))
    dates = _text(item.get("dates") or item.get("date") or item.get("period"))
    bullets = [_text(b) for b in _as_list(item.get("bullets") or item.get("details") or item.get("description"))]
    bullets = [b for b in bullets if b]
    return heading, subheading, dates, bullets


def render_markdown(data: dict[str, Any]) -> str:
    cand = _candidate(data)
    name = _text(cand.get("name") or "候选人")
    target_role = _text(cand.get("target_role") or cand.get("objective"))
    header = f"# {name}"
    if target_role:
        header += f"｜{target_role}"

    lines = [header, ""]
    contact = " ｜ ".join(part for part in _contact_parts(cand) if part)
    if contact:
        lines.extend([contact, ""])

    for section in _derive_sections(data):
        title = _text(section.get("title"))
        if not title:
            continue
        lines.extend([f"## {title}", ""])
        for item in _as_list(section.get("items")):
            heading, subheading, dates, bullets = _item_lines(item)
            row = " ｜ ".join(part for part in [heading, subheading, dates] if part)
            if row:
                lines.append(f"**{row}**")
            for bullet in bullets:
                lines.append(f"- {bullet}")
            if row:
                lines.append("")
        if lines and lines[-1] != "":
            lines.append("")
    return "\n".join(lines).rstrip() + "\n"


def _set_east_asia_font(run: Any, font_name: str) -> None:
    run.font.name = font_name
    try:
        from docx.oxml.ns import qn

        run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    except Exception:
        pass


def render_docx(data: dict[str, Any], path: Path) -> None:
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Cm, Pt
    except ImportError as exc:
        raise RuntimeError("python-docx is required to render DOCX") from exc

    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.35)
    section.left_margin = Cm(1.65)
    section.right_margin = Cm(1.65)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal.font.size = Pt(10)
    normal.paragraph_format.space_after = Pt(3)
    normal.paragraph_format.line_spacing = 1.08

    cand = _candidate(data)
    name = _text(cand.get("name") or "候选人")
    target_role = _text(cand.get("target_role") or cand.get("objective"))
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(name + (f"｜{target_role}" if target_role else ""))
    run.bold = True
    run.font.size = Pt(18)
    _set_east_asia_font(run, "Microsoft YaHei")

    contact = " ｜ ".join(part for part in _contact_parts(cand) if part)
    if contact:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(contact)
        r.font.size = Pt(9)
        _set_east_asia_font(r, "Microsoft YaHei")

    for section_data in _derive_sections(data):
        title_text = _text(section_data.get("title"))
        if not title_text:
            continue
        heading = doc.add_paragraph()
        heading.paragraph_format.space_before = Pt(7)
        heading.paragraph_format.space_after = Pt(2)
        r = heading.add_run(title_text)
        r.bold = True
        r.font.size = Pt(12)
        _set_east_asia_font(r, "Microsoft YaHei")

        for item in _as_list(section_data.get("items")):
            heading_text, subheading, dates, bullets = _item_lines(item)
            row = " ｜ ".join(part for part in [heading_text, subheading, dates] if part)
            if row:
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after = Pt(1)
                r = p.add_run(row)
                r.bold = True
                r.font.size = Pt(10)
                _set_east_asia_font(r, "Microsoft YaHei")
            for bullet in bullets:
                p = doc.add_paragraph(style=None)
                p.paragraph_format.left_indent = Cm(0.35)
                p.paragraph_format.first_line_indent = Cm(-0.18)
                p.paragraph_format.space_after = Pt(2)
                r = p.add_run("• " + bullet)
                r.font.size = Pt(9.5)
                _set_east_asia_font(r, "Microsoft YaHei")

    doc.save(path)


def main() -> int:
    parser = argparse.ArgumentParser(description="Render a Chinese resume from JSON.")
    parser.add_argument("input_json", type=Path)
    parser.add_argument("--output-dir", type=Path, default=Path("."))
    parser.add_argument("--basename", default="")
    parser.add_argument("--markdown-only", action="store_true")
    args = parser.parse_args()

    data = json.loads(args.input_json.read_text(encoding="utf-8"))
    cand = _candidate(data)
    basename = args.basename or "-".join(
        part for part in [_text(cand.get("name") or "候选人"), _text(cand.get("target_role") or "简历")] if part
    )
    args.output_dir.mkdir(parents=True, exist_ok=True)

    markdown_path = args.output_dir / f"{basename}-附件简历.md"
    markdown_path.write_text(render_markdown(data), encoding="utf-8")
    print(f"Wrote {markdown_path}")

    if not args.markdown_only:
        docx_path = args.output_dir / f"{basename}-附件简历.docx"
        render_docx(data, docx_path)
        print(f"Wrote {docx_path}")

    return 0


if __name__ == "__main__":
    raise SystemExit(main())
